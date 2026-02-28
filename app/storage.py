from __future__ import annotations

import base64
import re
import sqlite3
from pathlib import Path
from typing import Any, Optional

ROOT_DIR = Path(__file__).resolve().parents[1]
DB_DIR = ROOT_DIR / "data"
DB_PATH = DB_DIR / "ppt_mvp.db"


def _connect() -> sqlite3.Connection:
    DB_DIR.mkdir(parents=True, exist_ok=True)
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    try:
        # 若文件是无效 SQLite（如 CI 中未拉取 LFS 对象的 pointer 文本），自动重建。
        conn.execute("PRAGMA schema_version").fetchone()
        return conn
    except sqlite3.DatabaseError:
        conn.close()
        try:
            DB_PATH.unlink(missing_ok=True)
        except Exception:
            pass
        rebuilt = sqlite3.connect(DB_PATH)
        rebuilt.row_factory = sqlite3.Row
        return rebuilt


def init_db() -> None:
    with _connect() as conn:
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS import_jobs (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                source_filename TEXT NOT NULL,
                chapter_count INTEGER NOT NULL,
                zip_filename TEXT NOT NULL,
                zip_blob BLOB NOT NULL,
                created_at TEXT NOT NULL DEFAULT (datetime('now'))
            )
            """
        )
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS chapter_assets (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                job_id INTEGER NOT NULL,
                chapter_index INTEGER NOT NULL,
                title TEXT NOT NULL,
                summary TEXT NOT NULL,
                content TEXT NOT NULL,
                slide_count INTEGER NOT NULL,
                ppt_filename TEXT NOT NULL,
                ppt_blob BLOB NOT NULL,
                md_filename TEXT NOT NULL,
                md_text TEXT NOT NULL,
                word_filename TEXT,
                word_blob BLOB,
                created_at TEXT NOT NULL DEFAULT (datetime('now')),
                FOREIGN KEY(job_id) REFERENCES import_jobs(id)
            )
            """
        )
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS session_records (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                raw_query TEXT NOT NULL,
                generated_prompt TEXT NOT NULL,
                industry TEXT NOT NULL,
                customer TEXT NOT NULL,
                duration TEXT NOT NULL,
                product_name TEXT NOT NULL,
                visit_role TEXT NOT NULL,
                business_domains TEXT NOT NULL,
                created_at TEXT NOT NULL DEFAULT (datetime('now'))
            )
            """
        )
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS unified_templates (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                filename TEXT NOT NULL,
                ppt_blob BLOB NOT NULL,
                created_at TEXT NOT NULL DEFAULT (datetime('now'))
            )
            """
        )
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS app_settings (
                key TEXT PRIMARY KEY,
                value TEXT NOT NULL,
                updated_at TEXT NOT NULL DEFAULT (datetime('now'))
            )
            """
        )
        _ensure_column(conn, "chapter_assets", "word_filename", "TEXT")
        _ensure_column(conn, "chapter_assets", "word_blob", "BLOB")
        conn.commit()


def save_import_result(
    *,
    source_filename: str,
    chapters: list[dict[str, Any]],
    zip_bytes: bytes,
    zip_filename: str = "ppt_分拆章节.zip",
) -> dict[str, Any]:
    init_db()
    with _connect() as conn:
        cur = conn.execute(
            """
            INSERT INTO import_jobs (source_filename, chapter_count, zip_filename, zip_blob)
            VALUES (?, ?, ?, ?)
            """,
            (source_filename, len(chapters), zip_filename, sqlite3.Binary(zip_bytes)),
        )
        job_id = int(cur.lastrowid)

        for idx, chapter in enumerate(chapters, start=1):
            title = str(chapter.get("title") or f"章节{idx}")
            summary = str(chapter.get("summary") or "")
            content = str(chapter.get("content") or "")
            slide_count = int(chapter.get("slide_count") or 0)
            ppt_b64 = str(chapter.get("ppt_base64") or "")
            ppt_blob = base64.b64decode(ppt_b64) if ppt_b64 else b""
            safe_title = _safe_name(title)
            ppt_filename = f"章节{idx}_{safe_title}.pptx"
            md_filename = f"章节{idx}_{safe_title}.md"
            word_filename = f"章节{idx}_{safe_title}.docx"
            md_text = _build_md(title=title, summary=summary, content=content, slide_count=slide_count)
            word_blob = _build_word_bytes(title=title, summary=summary, content=content, slide_count=slide_count)

            conn.execute(
                """
                INSERT INTO chapter_assets (
                    job_id, chapter_index, title, summary, content, slide_count,
                    ppt_filename, ppt_blob, md_filename, md_text, word_filename, word_blob
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    job_id,
                    idx,
                    title,
                    summary,
                    content,
                    slide_count,
                    ppt_filename,
                    sqlite3.Binary(ppt_blob),
                    md_filename,
                    md_text,
                    word_filename,
                    sqlite3.Binary(word_blob),
                ),
            )

        conn.commit()

    return {
        "job_id": job_id,
        "chapter_count": len(chapters),
        "zip_filename": zip_filename,
    }


def get_db_info() -> dict[str, Any]:
    init_db()
    with _connect() as conn:
        table_names = [
            row["name"]
            for row in conn.execute(
                "SELECT name FROM sqlite_master WHERE type='table' AND name NOT LIKE 'sqlite_%' ORDER BY name"
            )
        ]
        tables: list[dict[str, Any]] = []
        for name in table_names:
            count = conn.execute(f"SELECT COUNT(*) AS c FROM {name}").fetchone()["c"]
            tables.append({"name": name, "rows": int(count)})

    return {
        "db_type": "SQLite",
        "database_name": DB_PATH.name,
        "database_path": str(DB_PATH),
        "tables": tables,
    }


def list_import_jobs(limit: int = 100) -> list[dict[str, Any]]:
    init_db()
    with _connect() as conn:
        rows = conn.execute(
            """
            SELECT id, source_filename, chapter_count, zip_filename, created_at
            FROM import_jobs
            ORDER BY id DESC
            LIMIT ?
            """,
            (max(1, min(limit, 500)),),
        ).fetchall()
    return [dict(row) for row in rows]


def get_job_detail(job_id: int) -> dict[str, Any] | None:
    init_db()
    with _connect() as conn:
        job = conn.execute(
            """
            SELECT id, source_filename, chapter_count, zip_filename, created_at
            FROM import_jobs
            WHERE id = ?
            """,
            (job_id,),
        ).fetchone()
        if not job:
            return None
        chapters = conn.execute(
            """
            SELECT id, chapter_index, title, summary, slide_count, ppt_filename, md_filename, word_filename, created_at
            FROM chapter_assets
            WHERE job_id = ?
            ORDER BY chapter_index ASC
            """,
            (job_id,),
        ).fetchall()
    return {"job": dict(job), "chapters": [dict(row) for row in chapters]}


def get_job_zip_blob(job_id: int) -> tuple[str, bytes] | None:
    init_db()
    with _connect() as conn:
        row = conn.execute("SELECT zip_filename, zip_blob FROM import_jobs WHERE id = ?", (job_id,)).fetchone()
    if not row:
        return None
    return str(row["zip_filename"]), bytes(row["zip_blob"] or b"")


def get_chapter_file_blob(chapter_id: int, file_type: str) -> tuple[str, bytes, str] | None:
    init_db()
    field_map = {
        "ppt": ("ppt_filename", "ppt_blob", "application/vnd.openxmlformats-officedocument.presentationml.presentation"),
        "md": ("md_filename", "md_text", "text/markdown; charset=utf-8"),
        "word": ("word_filename", "word_blob", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
    }
    if file_type not in field_map:
        return None
    filename_col, blob_col, media_type = field_map[file_type]
    with _connect() as conn:
        row = conn.execute(
            f"SELECT {filename_col} AS filename, {blob_col} AS payload FROM chapter_assets WHERE id = ?",
            (chapter_id,),
        ).fetchone()
    if not row:
        return None
    payload = row["payload"]
    if payload is None:
        return None
    if isinstance(payload, str):
        payload_bytes = payload.encode("utf-8")
    else:
        payload_bytes = bytes(payload)
    return str(row["filename"]), payload_bytes, media_type


def search_top_chapter_ppts(
    *,
    product_name: str,
    business_domains: list[str],
    visit_role: str,
    limit: int = 3,
) -> list[dict[str, Any]]:
    """
    按产品/业务域/拜访角色在章节内容中做关键词打分，返回前 N 个章节 PPT。
    """
    init_db()
    limit = max(1, min(limit, 20))
    product = (product_name or "").strip()
    role = (visit_role or "").strip()
    domains = [x.strip() for x in business_domains if x and x.strip()]

    with _connect() as conn:
        rows = conn.execute(
            """
            SELECT
                c.id,
                c.job_id,
                c.chapter_index,
                c.title,
                c.summary,
                c.content,
                c.ppt_filename,
                c.ppt_blob,
                c.created_at,
                j.source_filename
            FROM chapter_assets c
            LEFT JOIN import_jobs j ON j.id = c.job_id
            ORDER BY c.id DESC
            """
        ).fetchall()

    scored: list[tuple[int, int, dict[str, Any]]] = []
    for row in rows:
        text = " ".join(
            [
                str(row["title"] or ""),
                str(row["summary"] or ""),
                str(row["content"] or ""),
                str(row["source_filename"] or ""),
            ]
        )
        score = 0
        if product and product in text:
            score += 5
        if role and role in text:
            score += 2
        for domain in domains:
            if domain in text:
                score += 3
        item = {
            "chapter_id": int(row["id"]),
            "job_id": int(row["job_id"] or 0),
            "chapter_index": int(row["chapter_index"] or 0),
            "title": str(row["title"] or ""),
            "ppt_filename": str(row["ppt_filename"] or "matched.pptx"),
            "ppt_blob": bytes(row["ppt_blob"] or b""),
            "source_filename": str(row["source_filename"] or ""),
            "created_at": str(row["created_at"] or ""),
            "score": score,
        }
        scored.append((score, int(row["id"]), item))

    # 先取有命中的高分项，再用最新记录补齐到 limit
    scored.sort(key=lambda x: (x[0], x[1]), reverse=True)
    chosen: list[dict[str, Any]] = []
    used_ids: set[int] = set()
    for score, _, item in scored:
        if score <= 0:
            continue
        if item["chapter_id"] in used_ids:
            continue
        if not item["ppt_blob"]:
            continue
        used_ids.add(item["chapter_id"])
        chosen.append(item)
        if len(chosen) >= limit:
            return chosen

    for _, _, item in scored:
        if item["chapter_id"] in used_ids:
            continue
        if not item["ppt_blob"]:
            continue
        used_ids.add(item["chapter_id"])
        chosen.append(item)
        if len(chosen) >= limit:
            break
    return chosen


def save_session_record(
    *,
    raw_query: str,
    generated_prompt: str,
    industry: str,
    customer: str,
    duration: str,
    product_name: str,
    visit_role: str,
    business_domains: list[str],
) -> dict[str, Any]:
    init_db()
    domains_text = "、".join([x.strip() for x in business_domains if x and x.strip()])
    with _connect() as conn:
        cur = conn.execute(
            """
            INSERT INTO session_records (
                raw_query, generated_prompt, industry, customer, duration,
                product_name, visit_role, business_domains
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                raw_query.strip(),
                generated_prompt.strip(),
                industry.strip(),
                customer.strip(),
                duration.strip(),
                product_name.strip(),
                visit_role.strip(),
                domains_text,
            ),
        )
        conn.commit()
        record_id = int(cur.lastrowid)
        row = conn.execute(
            """
            SELECT id, raw_query, generated_prompt, industry, customer, duration,
                   product_name, visit_role, business_domains, created_at
            FROM session_records
            WHERE id = ?
            """,
            (record_id,),
        ).fetchone()
    return dict(row)


def list_session_records(limit: int = 100) -> list[dict[str, Any]]:
    init_db()
    with _connect() as conn:
        rows = conn.execute(
            """
            SELECT id, raw_query, generated_prompt, industry, customer, duration,
                   product_name, visit_role, business_domains, created_at
            FROM session_records
            ORDER BY id DESC
            LIMIT ?
            """,
            (max(1, min(limit, 500)),),
        ).fetchall()
    return [dict(row) for row in rows]


def save_unified_template(*, filename: str, ppt_bytes: bytes) -> dict[str, Any]:
    init_db()
    with _connect() as conn:
        cur = conn.execute(
            "INSERT INTO unified_templates (filename, ppt_blob) VALUES (?, ?)",
            (filename.strip() or "统一模板.pptx", sqlite3.Binary(ppt_bytes)),
        )
        template_id = int(cur.lastrowid)
        _set_active_template_id(conn, template_id)
        conn.commit()
        row = conn.execute(
            """
            SELECT id, filename, created_at, length(ppt_blob) AS file_size
            FROM unified_templates
            WHERE id = ?
            """,
            (template_id,),
        ).fetchone()
    return dict(row)


def get_latest_unified_template_meta() -> dict[str, Any] | None:
    init_db()
    with _connect() as conn:
        row = conn.execute(
            """
            SELECT id, filename, created_at, length(ppt_blob) AS file_size
            FROM unified_templates
            ORDER BY id DESC
            LIMIT 1
            """
        ).fetchone()
    return dict(row) if row else None


def get_latest_unified_template_blob() -> tuple[str, bytes] | None:
    init_db()
    with _connect() as conn:
        row = conn.execute(
            """
            SELECT filename, ppt_blob
            FROM unified_templates
            ORDER BY id DESC
            LIMIT 1
            """
        ).fetchone()
    if not row:
        return None
    return str(row["filename"]), bytes(row["ppt_blob"] or b"")


def list_unified_templates(limit: int = 200) -> list[dict[str, Any]]:
    init_db()
    with _connect() as conn:
        active_id = _get_active_template_id(conn)
        rows = conn.execute(
            """
            SELECT id, filename, created_at, length(ppt_blob) AS file_size
            FROM unified_templates
            ORDER BY id DESC
            LIMIT ?
            """,
            (max(1, min(limit, 1000)),),
        ).fetchall()
    # 同名且同大小的模板仅保留最新一条（按 id DESC 查询结果的首条）
    seen_keys: set[tuple[str, int]] = set()
    items: list[dict[str, Any]] = []
    for row in rows:
        item = dict(row)
        key = (str(item.get("filename") or ""), int(item.get("file_size") or 0))
        if key in seen_keys:
            continue
        seen_keys.add(key)
        items.append(item)
    for item in items:
        item["is_active"] = active_id is not None and int(item["id"]) == int(active_id)
    return items


def set_active_unified_template(template_id: int) -> dict[str, Any] | None:
    init_db()
    with _connect() as conn:
        row = conn.execute(
            """
            SELECT id, filename, created_at, length(ppt_blob) AS file_size
            FROM unified_templates
            WHERE id = ?
            """,
            (template_id,),
        ).fetchone()
        if not row:
            return None
        _set_active_template_id(conn, template_id)
        conn.commit()
    item = dict(row)
    item["is_active"] = True
    return item


def get_active_unified_template_meta() -> dict[str, Any] | None:
    init_db()
    with _connect() as conn:
        active_id = _get_active_template_id(conn)
        if active_id is not None:
            row = conn.execute(
                """
                SELECT id, filename, created_at, length(ppt_blob) AS file_size
                FROM unified_templates
                WHERE id = ?
                """,
                (active_id,),
            ).fetchone()
            if row:
                item = dict(row)
                item["is_active"] = True
                return item
        row = conn.execute(
            """
            SELECT id, filename, created_at, length(ppt_blob) AS file_size
            FROM unified_templates
            ORDER BY id DESC
            LIMIT 1
            """
        ).fetchone()
        if not row:
            return None
        fallback_id = int(row["id"])
        _set_active_template_id(conn, fallback_id)
        conn.commit()
        item = dict(row)
        item["is_active"] = True
        return item


def get_active_unified_template_blob() -> tuple[int, str, bytes] | None:
    init_db()
    with _connect() as conn:
        active_id = _get_active_template_id(conn)
        row = None
        if active_id is not None:
            row = conn.execute(
                "SELECT id, filename, ppt_blob FROM unified_templates WHERE id = ?",
                (active_id,),
            ).fetchone()
        if not row:
            row = conn.execute(
                "SELECT id, filename, ppt_blob FROM unified_templates ORDER BY id DESC LIMIT 1"
            ).fetchone()
            if row:
                _set_active_template_id(conn, int(row["id"]))
                conn.commit()
    if not row:
        return None
    return int(row["id"]), str(row["filename"]), bytes(row["ppt_blob"] or b"")


def get_unified_template_blob_by_id(template_id: int) -> tuple[str, bytes] | None:
    init_db()
    with _connect() as conn:
        row = conn.execute(
            "SELECT filename, ppt_blob FROM unified_templates WHERE id = ?",
            (template_id,),
        ).fetchone()
    if not row:
        return None
    return str(row["filename"]), bytes(row["ppt_blob"] or b"")


def delete_unified_template(template_id: int) -> dict[str, Any] | None:
    init_db()
    with _connect() as conn:
        row = conn.execute("SELECT id FROM unified_templates WHERE id = ?", (template_id,)).fetchone()
        if not row:
            return None
        conn.execute("DELETE FROM unified_templates WHERE id = ?", (template_id,))

        next_row = conn.execute(
            "SELECT id FROM unified_templates ORDER BY id DESC LIMIT 1"
        ).fetchone()
        if next_row:
            _set_active_template_id(conn, int(next_row["id"]))
        else:
            conn.execute("DELETE FROM app_settings WHERE key = 'active_unified_template_id'")
        conn.commit()
    return {"deleted_template_id": template_id, "next_active_template_id": int(next_row["id"]) if next_row else None}


def _get_active_template_id(conn: sqlite3.Connection) -> Optional[int]:
    row = conn.execute(
        "SELECT value FROM app_settings WHERE key = 'active_unified_template_id'"
    ).fetchone()
    if not row:
        return None
    try:
        return int(row["value"])
    except Exception:
        return None


def _set_active_template_id(conn: sqlite3.Connection, template_id: int) -> None:
    conn.execute(
        """
        INSERT INTO app_settings (key, value, updated_at)
        VALUES ('active_unified_template_id', ?, datetime('now'))
        ON CONFLICT(key) DO UPDATE SET value = excluded.value, updated_at = datetime('now')
        """,
        (str(template_id),),
    )


def delete_job(job_id: int) -> bool:
    """删除整件文件（job）及其所有章节内容。"""
    init_db()
    with _connect() as conn:
        exists = conn.execute("SELECT 1 FROM import_jobs WHERE id = ?", (job_id,)).fetchone()
        if not exists:
            return False
        conn.execute("DELETE FROM chapter_assets WHERE job_id = ?", (job_id,))
        conn.execute("DELETE FROM import_jobs WHERE id = ?", (job_id,))
        conn.commit()
    return True


def delete_chapter(chapter_id: int) -> dict[str, Any] | None:
    """删除单章节，并同步更新所属 job 的章节数和 ZIP。"""
    init_db()
    with _connect() as conn:
        row = conn.execute("SELECT id, job_id FROM chapter_assets WHERE id = ?", (chapter_id,)).fetchone()
        if not row:
            return None
        job_id = int(row["job_id"])
        conn.execute("DELETE FROM chapter_assets WHERE id = ?", (chapter_id,))
        _rebuild_job_zip_and_count(conn, job_id)
        conn.commit()

        left = conn.execute("SELECT chapter_count FROM import_jobs WHERE id = ?", (job_id,)).fetchone()
        chapter_count = int(left["chapter_count"]) if left else 0
    return {"job_id": job_id, "chapter_count": chapter_count}


def _safe_name(name: str) -> str:
    cleaned = re.sub(r"[^\w\-\u4e00-\u9fff]+", "_", name or "")
    cleaned = cleaned.strip("_ ").strip()
    return (cleaned or "未命名章节")[:40]


def _build_md(*, title: str, summary: str, content: str, slide_count: int) -> str:
    return (
        f"# {title}\n\n"
        f"- 页数: {slide_count}\n\n"
        f"## 概要总结\n\n{summary or '暂无摘要'}\n\n"
        f"## 章节正文\n\n{content or '暂无正文内容'}\n"
    )


def _build_word_bytes(*, title: str, summary: str, content: str, slide_count: int) -> bytes:
    try:
        from docx import Document
    except Exception:
        # python-docx 不可用时，回退为 markdown 文本内容
        return _build_md(title=title, summary=summary, content=content, slide_count=slide_count).encode("utf-8")

    doc = Document()
    doc.add_heading(_xml_safe_text(title) or "未命名章节", level=1)
    doc.add_paragraph(f"页数: {slide_count}")
    doc.add_heading("概要总结", level=2)
    doc.add_paragraph(_xml_safe_text(summary) or "暂无摘要")
    doc.add_heading("章节正文", level=2)
    doc.add_paragraph(_xml_safe_text(content) or "暂无正文内容")
    import io

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _ensure_column(conn: sqlite3.Connection, table_name: str, column_name: str, column_type: str) -> None:
    rows = conn.execute(f"PRAGMA table_info({table_name})").fetchall()
    existing = {row[1] for row in rows}
    if column_name not in existing:
        try:
            conn.execute(f"ALTER TABLE {table_name} ADD COLUMN {column_name} {column_type}")
        except sqlite3.OperationalError as exc:
            # 热重载并发启动时，可能出现重复加列竞态；重复列错误可安全忽略。
            if "duplicate column name" not in str(exc).lower():
                raise


def _xml_safe_text(value: str | None) -> str:
    text = value or ""
    # XML 1.0 不允许控制字符（除 \t \n \r）
    return "".join(ch for ch in text if ch == "\t" or ch == "\n" or ch == "\r" or ord(ch) >= 0x20)


def _rebuild_job_zip_and_count(conn: sqlite3.Connection, job_id: int) -> None:
    """基于当前章节记录重建 zip_blob 与 chapter_count。"""
    import io
    import zipfile

    chapters = conn.execute(
        """
        SELECT id, chapter_index, ppt_filename, ppt_blob, md_filename, md_text, word_filename, word_blob
        FROM chapter_assets
        WHERE job_id = ?
        ORDER BY chapter_index ASC, id ASC
        """,
        (job_id,),
    ).fetchall()

    # 重排章节序号，保持连续（1..N）
    for new_idx, ch in enumerate(chapters, start=1):
        if int(ch["chapter_index"]) != new_idx:
            conn.execute("UPDATE chapter_assets SET chapter_index = ? WHERE id = ?", (new_idx, int(ch["id"])))

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for ch in chapters:
            if ch["ppt_filename"] and ch["ppt_blob"] is not None:
                zf.writestr(str(ch["ppt_filename"]), bytes(ch["ppt_blob"]))
            if ch["md_filename"] and ch["md_text"] is not None:
                zf.writestr(str(ch["md_filename"]), str(ch["md_text"]).encode("utf-8"))
            if ch["word_filename"] and ch["word_blob"] is not None:
                zf.writestr(str(ch["word_filename"]), bytes(ch["word_blob"]))

    conn.execute(
        "UPDATE import_jobs SET chapter_count = ?, zip_blob = ? WHERE id = ?",
        (len(chapters), sqlite3.Binary(buf.getvalue()), job_id),
    )
